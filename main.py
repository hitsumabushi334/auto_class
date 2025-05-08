# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import ttk, messagebox
import threading
import time
from datetime import datetime
import os
import logging  # logging モジュールをインポート
import traceback  # スタックトレース取得のため

from certifi import contents
import win32gui  # ウィンドウ操作のため追加
import google.genai  # google モジュールを明示的にインポート
from google import genai
from docx import Document  # mdファイル生成のため追加
from docx.shared import Inches  # mdファイル生成のため追加 (必要に応じて)
import json
from markd import Markdown

# import sounddevice as sd # sounddevice を削除
import queue  # 音声データキューのため追加

# import pyaudio  # PyAudio を削除
import soundcard as sc  # SoundCard を追加

# import ffmpeg  # ffmpeg-python は削除
import moviepy.editor as mpe  # MoviePy を追加
from moviepy.audio.AudioClip import AudioArrayClip  # AudioArrayClip を直接インポート

import mss  # 画面キャプチャのため追加
import cv2
from PIL import Image, ImageGrab, UnidentifiedImageError
import numpy as np
from dotenv import load_dotenv

# --- ロギング設定 ---
log_filename = "slide_capture_app.log"
log_format = "%(asctime)s - %(levelname)s - %(threadName)s - %(message)s"
logging.basicConfig(
    level=logging.INFO,
    format=log_format,
    handlers=[
        logging.FileHandler(log_filename, encoding="utf-8"),  # ファイル出力
        logging.StreamHandler(),  # コンソールにも出力 (デバッグ用)
    ],
)
logger = logging.getLogger(__name__)


class SlideCaptureApp:
    def __init__(self, root):
        self.root = root
        self.root.title("スライドキャプチャ＆録画")
        # UIの高さを増やして新しい要素を配置
        self.root.geometry("600x550")  # サイズ調整
        self.root.update_idletasks()
        self.root.minsize(600, 550)

        # --- 状態変数 ---
        self.is_capturing_screenshot = False  # スクリーンショット中フラグ
        self.is_recording = False  # 録画中フラグ
        self.screenshot_thread = None
        self.recording_thread = None
        self.audio_recording_thread = None  # 音声録音スレッド
        self.recorded_frames = []  # 録画フレームを一時保存するリスト
        self.audio_queue = queue.Queue()  # 音声データを一時保存するキュー
        self.last_screenshot_image = None
        self.screenshot_saved_count = 0
        self.last_saved_screenshot_filename = ""
        self.recording_start_time = None
        self.last_saved_video_filename = ""
        self.save_folder_name = tk.StringVar()
        self.selected_window_handle = tk.IntVar(value=0)  # 選択されたウィンドウハンドル
        self.error_occurred_in_screenshot_thread = False
        self.error_occurred_in_recording_thread = False
        self.note_creation_status = tk.StringVar(value="")  # ノート作成ステータス
        self.note_result_message = tk.StringVar(value="")  # ノート作成結果
        self.gemini_client = None  # Gemini API クライアント
        self.audio_sample_rate = None  # SoundCard で取得したサンプルレートを保存
        self.audio_channels = None  # SoundCard で取得したチャンネル数を保存
        self.last_sound_time = None  # 最後に音声を検知した時刻
        self.no_sound_timeout_seconds = 180  # 無音状態のタイムアウト秒数 (3分)
        self.silence_threshold = (
            0.01  # 無音と判定する振幅の閾値 (0.0 から 1.0 の範囲で調整)
        )
        # ユーザー指示に基づきモデル名を更新 (ただし、実際のAPI呼び出しでは利用可能なモデルを確認すること)
        self.gemini_model_options = [
            "gemini-2.5-flash-preview-04-17",  # 短時間用 (デフォルト)
            "gemini-2.0-flash",  # 長時間用
        ]

        self.selected_gemini_model = tk.StringVar(
            value=self.gemini_model_options[0]  # デフォルトは短時間用
        )

        # --- Gemini API 設定 ---
        try:
            load_dotenv()  # .env ファイルを読み込む
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key:
                logger.warning(
                    "環境変数 GEMINI_API_KEY が設定されていません。ノート作成機能は利用できません。"
                )
                messagebox.showwarning(
                    "APIキー未設定",
                    "環境変数 GEMINI_API_KEY が設定されていません。\nノート作成機能は利用できません。",
                )
            else:
                # 動画を扱えるモデルを指定 (例: gemini-1.5-pro-latest)
                # 注意: モデル名や利用可否は変更される可能性があるため、ドキュメントを確認すること
                # self.gemini_model = genai.GenerativeModel('gemini-1.5-pro-latest') # 後で使う
                logger.info("Gemini API クライアントを設定しました。")
                self.gemini_client = genai.Client(api_key=api_key)
        except Exception as e:
            logger.exception("Gemini API の設定中にエラーが発生しました。")
            messagebox.showerror(
                "API設定エラー", f"Gemini API の設定に失敗しました:\n{e}"
            )

        # --- UI要素の作成 ---

        # --- 保存フォルダ設定 ---
        folder_frame = ttk.Frame(root, padding="10")
        folder_frame.pack(fill=tk.X)
        folder_label = ttk.Label(folder_frame, text="保存フォルダ名:")
        folder_label.pack(side=tk.LEFT, padx=(0, 5))
        self.folder_entry = ttk.Entry(
            folder_frame, textvariable=self.save_folder_name, width=45  # 幅調整
        )
        self.folder_entry.pack(side=tk.LEFT, expand=True, fill=tk.X)

        # --- ウィンドウ選択 ---
        window_selection_frame = ttk.LabelFrame(
            root, text="録画対象ウィンドウ選択", padding="10"
        )
        window_selection_frame.pack(fill=tk.X, padx=10, pady=5)

        self.window_listbox = tk.Listbox(
            window_selection_frame, height=5, exportselection=False
        )
        self.window_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.window_listbox.bind("<<ListboxSelect>>", self.on_window_select)  # 実装

        window_scrollbar = ttk.Scrollbar(
            window_selection_frame,
            orient=tk.VERTICAL,
            command=self.window_listbox.yview,
        )
        window_scrollbar.pack(side=tk.LEFT, fill=tk.Y)
        self.window_listbox.config(yscrollcommand=window_scrollbar.set)

        self.refresh_window_list_button = ttk.Button(
            window_selection_frame,
            text="更新",
            command=self.refresh_window_list,  # あとで実装
        )
        self.refresh_window_list_button.pack(side=tk.LEFT, padx=5)

        # --- Gemini モデル選択 ---
        model_selection_frame = ttk.LabelFrame(
            root, text="Gemini モデル選択", padding="10"
        )
        model_selection_frame.pack(fill=tk.X, padx=10, pady=5)

        model_label = ttk.Label(model_selection_frame, text="モデル:")
        model_label.pack(side=tk.LEFT, padx=(0, 5))

        self.model_combobox = ttk.Combobox(
            model_selection_frame,
            textvariable=self.selected_gemini_model,
            values=self.gemini_model_options,
            state="readonly",  # ユーザー入力不可
            width=40,  # 幅調整
        )
        self.model_combobox.pack(side=tk.LEFT, padx=(0, 5))  # 右に少し余白
        self.model_combobox.bind(
            "<<ComboboxSelected>>", self._on_gemini_model_selected
        )  # イベントハンドラをバインド

        # 用途表示ラベルを追加
        self.gemini_model_description_label = ttk.Label(model_selection_frame, text="")
        self.gemini_model_description_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # --- 操作ボタン (統合) ---
        button_frame = ttk.Frame(root, padding="10")
        button_frame.pack(fill=tk.X)

        self.start_button = ttk.Button(
            button_frame,
            text="開始",
            command=self.start_tasks,  # 新しい統合開始メソッド
        )
        self.start_button.pack(side=tk.LEFT, padx=5)
        self.stop_button = ttk.Button(
            button_frame,
            text="停止",
            command=self.stop_all_tasks,
            state=tk.DISABLED,  # 統合停止メソッド
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)

        # --- ステータス表示 ---
        status_frame = ttk.Frame(root, padding="10")
        status_frame.pack(fill=tk.BOTH, expand=True)

        # スクリーンショットステータス
        screenshot_status_frame = ttk.LabelFrame(
            status_frame, text="スクリーンショット", padding="5"
        )
        screenshot_status_frame.pack(fill=tk.X, pady=(0, 5))
        self.screenshot_status_label = ttk.Label(
            screenshot_status_frame,
            text="待機中...",
            anchor=tk.W,
            justify=tk.LEFT,
            wraplength=400,
        )
        self.screenshot_status_label.pack(fill=tk.X)

        # 録画ステータス
        recording_status_frame = ttk.LabelFrame(status_frame, text="録画", padding="5")
        recording_status_frame.pack(fill=tk.X, pady=(0, 5))
        self.recording_status_label = ttk.Label(
            recording_status_frame,
            text="待機中...",
            anchor=tk.W,
            justify=tk.LEFT,
            wraplength=400,
        )
        self.recording_status_label.pack(fill=tk.X)

        # ノート作成ステータス
        note_status_frame = ttk.LabelFrame(status_frame, text="ノート作成", padding="5")
        note_status_frame.pack(fill=tk.X, pady=(0, 5))
        self.note_status_label = ttk.Label(
            note_status_frame,
            textvariable=self.note_creation_status,
            anchor=tk.W,
            justify=tk.LEFT,
            wraplength=400,
        )
        self.note_status_label.pack(fill=tk.X)
        self.note_result_label = ttk.Label(
            note_status_frame,
            textvariable=self.note_result_message,
            anchor=tk.W,
            justify=tk.LEFT,
            wraplength=400,
            foreground="blue",
        )
        self.note_result_label.pack(fill=tk.X)

        # --- イベントバインド ---
        self.root.bind("<Escape>", lambda e: self.stop_all_tasks())  # Escで全停止
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.original_on_closing = self.on_closing  # 元の閉じる処理を保持

        logger.info("アプリケーションを初期化しました。")
        self.refresh_window_list()  # 初期ウィンドウリスト表示
        self._on_gemini_model_selected()  # 初期モデルの用途を表示

    # --- ウィンドウ選択関連メソッド ---
    def refresh_window_list(self):
        """実行中のウィンドウリストを取得し、リストボックスを更新する"""
        logger.info("ウィンドウリストを更新します。")
        self.window_listbox.delete(0, tk.END)
        self.window_handles = {}  # タイトルとハンドルのマッピングを保持

        try:

            def enum_windows_proc(hwnd, lParam):
                if win32gui.IsWindowVisible(hwnd):
                    text = win32gui.GetWindowText(hwnd)
                    if text:  # タイトルがあるウィンドウのみ
                        # 同じタイトルのウィンドウがある場合、ハンドルを追記して区別
                        display_text = f"{text} (HWND: {hwnd})"
                        self.window_listbox.insert(tk.END, display_text)
                        self.window_handles[display_text] = hwnd
                return True  # 列挙を続ける

            win32gui.EnumWindows(enum_windows_proc, None)
            logger.info(
                f"{self.window_listbox.size()} 個のウィンドウが見つかりました。"
            )
        except Exception as e:
            logger.exception("ウィンドウリストの取得中にエラーが発生しました。")
            messagebox.showerror(
                "エラー", f"ウィンドウリストの取得に失敗しました:\n{e}"
            )

        # self.start_recording_button.config(state=tk.DISABLED) # 個別ボタン削除

    def on_window_select(self, event):
        """リストボックスでウィンドウが選択されたときの処理"""
        selected_indices = self.window_listbox.curselection()
        if not selected_indices:
            self.selected_window_handle.set(0)
            # self.start_recording_button.config(state=tk.DISABLED) # 個別ボタン削除
            logger.info("ウィンドウ選択が解除されました。")
            return

        selected_index = selected_indices[0]
        selected_text = self.window_listbox.get(selected_index)

        # ハンドルを取得
        hwnd = self.window_handles.get(selected_text)
        if hwnd:
            self.selected_window_handle.set(hwnd)
            # self.start_recording_button.config(state=tk.NORMAL) # 個別ボタン削除
            logger.info(f"ウィンドウが選択されました: {selected_text} (HWND: {hwnd})")
            # 開始ボタンの状態は start_tasks で制御
        else:
            logger.error(
                f"選択されたテキストに対応するハンドルが見つかりません: {selected_text}"
            )
            self.selected_window_handle.set(0)
            # self.start_recording_button.config(state=tk.DISABLED) # 個別ボタン削除

    def _on_gemini_model_selected(self, event=None):
        """Geminiモデル選択コンボボックスの値が変更されたときに呼び出される"""
        selected_model = self.selected_gemini_model.get()
        description = ""
        if selected_model == "gemini-2.0-flash":
            description = "用途: 長時間動画向け"
        elif selected_model == "gemini-2.5-flash-preview-04-17":
            description = "用途: 短時間動画向け"
        else:
            # 予期しないモデルが選択された場合 (念のため)
            description = "用途: 不明"
            logger.warning(f"不明なGeminiモデルが選択されました: {selected_model}")

        self.gemini_model_description_label.config(text=description)
        logger.info(
            f"Geminiモデル '{selected_model}' が選択されました。用途: {description}"
        )

    # --- 統合開始・停止メソッド ---
    def start_tasks(self):
        """スクリーンショットと録画を開始する（ウィンドウ選択状態による）"""
        if self.is_capturing_screenshot or self.is_recording:
            logger.warning("既にタスクが実行中です。")
            return

        hwnd = self.selected_window_handle.get()

        # フォルダ設定は共通で最初に行う
        if not self.prepare_save_folder():
            return

        # スクリーンショットは常に開始
        self.start_screenshot_capture()

        # ウィンドウが選択されていれば録画も開始
        if hwnd != 0:
            self.start_recording(hwnd)  # 引数でハンドルを渡す
        else:
            logger.info(
                "ウィンドウが選択されていないため、スクリーンショットのみ開始します。"
            )

        # 統合ボタンの状態更新
        self.start_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        self.folder_entry.config(state=tk.DISABLED)
        self.refresh_window_list_button.config(
            state=tk.DISABLED
        )  # タスク実行中は更新不可
        self.window_listbox.config(state=tk.DISABLED)
        self.model_combobox.config(state=tk.DISABLED)  # モデル選択も無効化

    # --- 録画関連メソッド (修正) ---
    def start_recording(self, hwnd):  # 引数 hwnd を追加
        """指定されたウィンドウハンドルで録画を開始する"""
        # hwnd = self.selected_window_handle.get() # start_tasks から渡される
        if hwnd == 0:  # ハンドルが無効な場合は何もしない（基本的には呼ばれないはず）
            logger.error("録画開始エラー: 無効なウィンドウハンドルです。")
            return

        if self.is_recording:  # 既に録画中の場合は何もしない
            return

        logger.info(f"録画を開始します。対象ウィンドウハンドル: {hwnd}")
        # 録画スレッドを開始
        self.recorded_frames = []  # フレームリストを初期化
        self.recording_thread = threading.Thread(
            target=self.recording_loop,
            args=(hwnd,),
            name="RecordingThread",
            daemon=True,
        )
        self.is_recording = True  # is_recording は先に True にする
        self.recording_thread.start()

        # 音声録音スレッドを開始
        self.audio_queue = queue.Queue()  # キューを初期化
        self.audio_recording_thread = threading.Thread(
            target=self.audio_recording_loop, name="AudioRecordingThread", daemon=True
        )
        self.audio_recording_thread.start()

        # self.start_recording_button.config(state=tk.DISABLED) # 個別ボタン削除
        # self.stop_recording_button.config(state=tk.NORMAL) # 個別ボタン削除
        # ボタン状態は start_tasks / stop_all_tasks で制御
        self.recording_start_time = time.time()
        self.last_sound_time = time.time()  # 録画開始時は音があったとみなす
        self.update_recording_status()  # ステータス更新開始

    def recording_loop(self, hwnd):
        """指定されたウィンドウのフレームを1FPSで録画するループ処理"""
        logger.info(f"録画ループを開始します。対象ウィンドウハンドル: {hwnd}")

        try:
            # mssのスクリーンショット用オブジェクト
            sct = mss.mss()

            # ウィンドウの位置とサイズを取得
            try:
                window_rect = win32gui.GetWindowRect(hwnd)
                left, top, right, bottom = window_rect
                width = right - left
                height = bottom - top
                logger.info(
                    f"対象ウィンドウの位置とサイズ: ({left}, {top}, {width}, {height})"
                )

                # ウィンドウ領域の定義
                monitor = {"left": left, "top": top, "width": width, "height": height}
            except Exception as e:
                logger.exception(f"ウィンドウ情報の取得に失敗しました: {e}")
                self.error_occurred_in_recording_thread = True
                return

            # 録画開始時刻を記録
            start_time = time.time()
            frame_count = 0

            while self.is_recording:
                # --- ウィンドウ存在チェック ---
                if not win32gui.IsWindow(hwnd):
                    logger.warning(
                        f"録画対象ウィンドウ (HWND: {hwnd}) が見つかりません。録画を停止します。"
                    )
                    # UIスレッドで停止処理を呼び出す
                    self.root.after(0, self.stop_all_tasks)
                    break  # ループを抜ける
                # --- ここまで追加 ---

                try:
                    # 現在のフレーム時刻
                    current_time = time.time()
                    elapsed_time = current_time - start_time

                    # スクリーンショットを取得
                    screenshot = sct.grab(monitor)

                    # Pillowイメージに変換
                    img = Image.frombytes("RGB", screenshot.size, screenshot.rgb)

                    # OpenCV形式に変換
                    frame = np.array(img)
                    frame = cv2.cvtColor(frame, cv2.COLOR_RGB2BGR)

                    # フレームを保存
                    if frame is not None and frame.size > 0:
                        self.recorded_frames.append((frame, elapsed_time))
                        frame_count += 1
                        if frame_count % 10 == 0:  # 10フレームごとにログ出力
                            logger.info(
                                f"録画フレーム数: {frame_count}, 経過時間: {elapsed_time:.2f}秒"
                            )
                    else:
                        logger.warning("無効なフレームがスキップされました")

                    # 次のフレームタイミングまで待機（1FPS）
                    next_frame_time = start_time + (frame_count * 1.0)  # 1秒間隔
                    sleep_time = max(0, next_frame_time - time.time())
                    if sleep_time > 0:
                        time.sleep(sleep_time)

                except Exception as e:
                    logger.exception(f"フレーム取得中にエラーが発生しました: {e}")
                    self.error_occurred_in_recording_thread = True
                    time.sleep(1.0)  # エラー時は少し待機してリトライ

        except Exception as e:
            logger.exception(f"録画ループでエラーが発生しました: {e}")
            self.error_occurred_in_recording_thread = True

        finally:
            logger.info(
                f"録画ループを終了します。合計フレーム数: {len(self.recorded_frames)}"
            )

    def audio_recording_loop(self):
        """SoundCardを使用してシステムサウンド (ループバック) を録音するループ"""
        logger.info("音声録音ループ (SoundCard) を開始します")
        num_frames = 1024  # 一度に読み取るサンプル数 (SoundCardの推奨値に合わせる)

        try:
            # デフォルトのループバックデバイスを取得
            # SoundCard はデフォルトでスピーカーのループバックを選択しようとします
            microphone_device = sc.get_microphone(
                id=str(sc.default_speaker().name), include_loopback=True
            )
            logger.info(
                f"取得したマイクデバイス: {microphone_device.name}"
            )  # マイクデバイス名を確認
            logger.info(f"マイクデバイスの型: {type(microphone_device)}")

            # サンプルレートとチャンネル数をマイクデバイスから取得
            try:
                self.audio_sample_rate = microphone_device.samplerate
                self.audio_channels = microphone_device.channels
                logger.info(
                    f"デバイス情報 - サンプルレート: {self.audio_sample_rate}, チャンネル数: {self.audio_channels}"
                )
            except AttributeError as e:
                logger.error(
                    f"マイクデバイスからサンプルレートまたはチャンネル数を取得できませんでした: {e}"
                )
                # デフォルト値を設定するか、エラー処理を行う
                self.audio_sample_rate = 48000  # デフォルト値
                self.audio_channels = 2  # デフォルト値
                logger.warning(
                    f"デフォルト値を使用します - サンプルレート: {self.audio_sample_rate}, チャンネル数: {self.audio_channels}"
                )

            with microphone_device.recorder(
                samplerate=self.audio_sample_rate,  # デバイスから取得した値を使用
                channels=self.audio_channels,  # デバイスから取得した値を使用
                blocksize=num_frames,
            ) as mic:
                logger.info(
                    f"レコーダーオブジェクトの型: {type(mic)}"
                )  # mic の型を確認
                logger.info(
                    f"録音中: {microphone_device.name} (ループバック)"
                )  # マイクデバイス名を使用

                # 録音ループ
                while self.is_recording:
                    try:
                        # recordメソッドで指定したサンプル数を読み取る
                        data = mic.record(numframes=num_frames)
                        if data is not None and data.size > 0:
                            # --- 無音状態チェック ---
                            max_amplitude = np.max(np.abs(data))
                            current_time = time.time()
                            if max_amplitude < self.silence_threshold:
                                # 無音状態
                                if self.last_sound_time is not None:
                                    no_sound_duration = (
                                        current_time - self.last_sound_time
                                    )
                                    if (
                                        no_sound_duration
                                        > self.no_sound_timeout_seconds
                                    ):
                                        logger.info(
                                            f"{self.no_sound_timeout_seconds} 秒以上無音状態が続いたため、録画を停止します。"
                                        )
                                        self.root.after(0, self.stop_all_tasks)
                                        break  # ループを抜ける
                            else:
                                # 音声あり
                                self.last_sound_time = current_time
                            # --- ここまで追加 ---

                            # SoundCard は float32 の NumPy 配列を返す
                            self.audio_queue.put(data)
                        else:
                            logger.warning(
                                "SoundCard から None または空のデータが返されました。"
                            )
                            time.sleep(0.01)  # 少し待機
                    except Exception as e:
                        logger.exception(
                            f"音声データ読み取り/キュー追加中にエラー: {e}"
                        )
                        self.error_occurred_in_recording_thread = True
                        break  # ループ中断

        except Exception as e:
            logger.exception(f"音声録音 (SoundCard) 中にエラーが発生しました: {e}")
            messagebox.showerror(
                "録音エラー",
                f"音声録音の初期化または実行中にエラーが発生しました:\n{e}",
            )
            self.error_occurred_in_recording_thread = True

        finally:
            logger.info("音声録音ループ (SoundCard) を終了します")

    def _save_video_with_audio(self, output_filepath):
        """録画されたフレームと音声データからMoviePyを使って動画ファイルを生成・保存し、成功したらノート作成をトリガーする"""
        if not self.recorded_frames:
            logger.error("保存するフレームがありません")
            return False

        # 音声データが存在するか確認し、連結する
        audio_data = []
        if not self.audio_queue.empty():
            while not self.audio_queue.empty():
                audio_data.append(self.audio_queue.get())

        audio_array = None
        if audio_data:
            try:
                audio_array = np.concatenate(audio_data)
                logger.info(
                    f"音声データを連結しました。サンプル数: {len(audio_array)}, dtype: {audio_array.dtype}"
                )
            except ValueError as e:
                logger.error(f"音声データの連結に失敗しました: {e}")
                audio_array = None  # エラー時は音声なしとする

        try:
            # 保存先のディレクトリを確認
            output_dir = os.path.dirname(output_filepath)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            # フレームの情報を取得 (BGR -> RGBに変換)
            frames_rgb = []
            timestamps = []
            first_frame_shape = None
            for frame_bgr, ts in self.recorded_frames:
                if frame_bgr is not None and frame_bgr.size > 0:
                    try:
                        frame_rgb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)
                        if first_frame_shape is None:
                            first_frame_shape = frame_rgb.shape
                        # フレームサイズが異なる場合はリサイズ (最初のフレームに合わせる)
                        if frame_rgb.shape != first_frame_shape:
                            logger.warning(
                                f"フレームサイズが異なります。リサイズします: {frame_rgb.shape} -> {first_frame_shape}"
                            )
                            frame_rgb = cv2.resize(
                                frame_rgb, (first_frame_shape[1], first_frame_shape[0])
                            )
                        frames_rgb.append(frame_rgb)
                        timestamps.append(ts)
                    except cv2.error as e:
                        logger.error(
                            f"フレームのRGB変換中にエラー: {e}. スキップします。"
                        )
                else:
                    logger.warning("None または空のフレームをスキップしました。")

            if not frames_rgb:
                logger.error("有効なフレームがありませんでした。動画を保存できません。")
                return False

            logger.info(f"合計 {len(frames_rgb)} フレームを動画クリップに使用します。")

            # フレーム間の時間を計算してFPSを推定 (MoviePyは可変FPSを直接扱えないため)
            if len(timestamps) > 1:
                avg_interval = np.mean(np.diff(timestamps))
                fps = 1.0 / avg_interval if avg_interval > 0 else 1.0
                logger.info(f"推定FPS: {fps:.2f}")
            else:
                fps = 1.0  # フレームが1つしかない場合
                logger.warning("フレームが1つしかないため、FPSを1.0に設定します。")

            # MoviePyのVideoClipを作成
            video_clip = mpe.ImageSequenceClip(frames_rgb, fps=fps)

            # 音声データがある場合、AudioClipを作成して結合
            audio_clip = None
            final_clip = video_clip  # デフォルトは音声なし
            if audio_array is not None and self.audio_sample_rate is not None:
                try:
                    # MoviePyは通常、[-1, 1]の範囲のfloatを期待する
                    # SoundCardが返すデータ形式を確認し、必要なら正規化
                    if audio_array.dtype != np.float32:
                        logger.warning(
                            f"音声データのdtypeがfloat32ではありません: {audio_array.dtype}。変換を試みます。"
                        )
                        # 必要に応じて型変換や正規化を行う (例: int16 -> float32)
                        # この例では float32 を想定
                        pass  # 必要ならここに変換処理を追加

                    # チャンネル数が1の場合、ステレオに変換 (MoviePyがステレオを期待する場合がある)
                    # if self.audio_channels == 1 and audio_array.ndim == 1:
                    #      audio_array = np.column_stack((audio_array, audio_array))
                    #      logger.info("モノラル音声をステレオに変換しました。")

                    # 音声配列の形状を確認・整形
                    if audio_array.ndim == 1 and self.audio_channels == 2:
                        # モノラルデータが連結されて1次元になっている場合、2チャンネルに整形
                        logger.warning("1次元の音声配列を2チャンネルに整形します。")
                        try:
                            audio_array = audio_array.reshape(-1, self.audio_channels)
                        except ValueError as reshape_err:
                            logger.error(
                                f"音声配列の整形に失敗しました: {reshape_err}。音声なしで続行します。"
                            )
                            audio_array = None
                    elif audio_array.ndim == 1 and self.audio_channels == 1:
                        # 1チャンネルの場合はそのままで良いことが多い
                        pass
                    elif (
                        audio_array.ndim == 2
                        and audio_array.shape[1] == self.audio_channels
                    ):
                        # 既に正しい形式
                        pass
                    else:
                        logger.error(
                            f"音声配列の形状 ({audio_array.shape}) がチャンネル数 ({self.audio_channels}) と一致しません。音声なしで続行します。"
                        )
                        audio_array = None  # 不正な場合は音声なしに

                    if audio_array is not None:
                        try:
                            # AudioArrayClip を直接使用 (mpe. ではなく)
                            audio_clip = AudioArrayClip(
                                audio_array, fps=self.audio_sample_rate
                            )
                            # 動画の長さに合わせて音声クリップの長さを調整
                            if audio_clip.duration > video_clip.duration:
                                audio_clip = audio_clip.subclip(0, video_clip.duration)
                            elif audio_clip.duration < video_clip.duration:
                                # 必要に応じて無音を追加するか、動画を短くする
                                logger.warning(
                                    f"音声クリップ ({audio_clip.duration:.2f}s) が動画クリップ ({video_clip.duration:.2f}s) より短いです。"
                                )
                                # audio_clip = audio_clip.set_duration(video_clip.duration) # 最後のフレームを繰り返す場合
                                pass  # そのまま結合する

                            logger.info(
                                f"AudioArrayClipを作成しました。Duration: {audio_clip.duration:.2f}秒"
                            )
                            final_clip = video_clip.set_audio(audio_clip)
                            logger.info("動画と音声を結合しました。")
                        except Exception as audio_clip_err:
                            logger.exception(
                                f"AudioArrayClipの作成または結合中にエラー: {audio_clip_err}"
                            )
                            final_clip = video_clip  # エラー時は音声なし
                    else:
                        final_clip = video_clip  # 整形失敗などで音声なしになった場合
                        logger.info("音声データが不正なため、動画のみ保存します。")

                except Exception as e:
                    logger.exception(f"音声処理中に予期せぬエラーが発生しました: {e}")
                    final_clip = video_clip  # エラー時は音声なしで保存
            else:
                # final_clip = video_clip # audio_array や sample_rate がない場合 (既に上で設定済み)
                logger.info(
                    "音声データまたはサンプルレートがないため、動画のみ保存します。"
                )

            # 動画ファイルを書き出し
            logger.info(f"動画ファイルを書き出します: {output_filepath}")
            try:
                # codec='libx264' を指定して互換性を高める
                # audio_codec='aac' を指定 (多くのプレイヤーでサポート)
                # threads を指定してエンコードを高速化 (CPUコア数など)
                # preset='medium' などで品質と速度のバランスを取る
                final_clip.write_videofile(
                    output_filepath,
                    codec="libx264",
                    audio_codec="aac",
                    temp_audiofile="temp-audio.m4a",  # 一時ファイル名を指定
                    remove_temp=True,  # 一時ファイルを削除
                    threads=8,  # CPUコア数に合わせて調整
                    preset="medium",  # 品質と速度のバランス
                    logger=None,
                    ffmpeg_params=[
                        "-profile:v",
                        "baseline",  # H.264 Baseline Profile: 幅広い互換性
                        "-level",
                        "3.0",  # レベル3.0: 多くのデバイスでサポート
                        "-pix_fmt",
                        "yuv420p",  # ピクセルフォーマット: 最も一般的な形式
                        "-vf",
                        "scale=trunc(iw/2)*2:trunc(ih/2)*2",  # 解像度を偶数に調整 (互換性向上)
                    ],  # MoviePyのログを無効化 (Pythonのloggingを使用するため)
                )
                logger.info(f"動画ファイルを保存しました: {output_filepath}")
                # self.last_saved_video_filename = output_filepath # ここでは設定しない

                # ★★★ 保存成功後にノート作成をトリガー ★★★
                # UIスレッド経由で start_note_creation を呼び出す
                logger.info(
                    f"動画保存成功、ノート作成をトリガーします: {output_filepath}"
                )
                self.root.after(0, self.start_note_creation, output_filepath)
                return True  # 保存自体は成功
            except Exception as e:
                # ffmpeg のエラーメッセージを取得しようとする試み
                ffmpeg_error = ""
                if hasattr(e, "stderr"):
                    try:
                        ffmpeg_error = e.stderr.decode("utf-8", errors="ignore")
                        logger.error(f"FFmpegエラー出力:\n{ffmpeg_error}")
                    except Exception as decode_err:
                        logger.error(f"FFmpegエラー出力のデコードに失敗: {decode_err}")
                logger.exception(
                    f"MoviePyによる動画書き出し中にエラーが発生しました: {e}"
                )
                messagebox.showerror(
                    "動画保存エラー",
                    f"動画の保存に失敗しました:\n{e}\n\nFFmpegエラー:\n{ffmpeg_error[:500]}...",
                )  # エラーメッセージを短縮
                return False

        except Exception as e:
            logger.exception(f"動画保存処理全体でエラーが発生しました: {e}")
            messagebox.showerror(
                "動画保存エラー", f"動画の保存中に予期せぬエラーが発生しました:\n{e}"
            )
            return False
        finally:
            # クリップオブジェクトを閉じる (メモリ解放)
            if "video_clip" in locals() and video_clip:
                video_clip.close()
            if "audio_clip" in locals() and audio_clip:
                audio_clip.close()
            if "final_clip" in locals() and final_clip:
                final_clip.close()
            logger.info("動画保存処理を終了します。")

    def stop_recording(self):
        """録画と音声録音を停止し、動画ファイルを保存する"""
        if not self.is_recording:
            return

        logger.info("録画停止処理を開始します。")
        self.is_recording = False  # まずフラグを False にする

        # 録画スレッドの終了を待つ
        if self.recording_thread and self.recording_thread.is_alive():
            logger.info("録画スレッドの終了を待機します...")
            self.recording_thread.join(timeout=5.0)  # タイムアウトを設定
            if self.recording_thread.is_alive():
                logger.warning("録画スレッドが時間内に終了しませんでした。")
            else:
                logger.info("録画スレッドが終了しました。")
        self.recording_thread = None

        # 音声録音スレッドの終了を待つ
        if self.audio_recording_thread and self.audio_recording_thread.is_alive():
            logger.info("音声録音スレッドの終了を待機します...")
            self.audio_recording_thread.join(timeout=5.0)  # タイムアウトを設定
            if self.audio_recording_thread.is_alive():
                logger.warning("音声録音スレッドが時間内に終了しませんでした。")
            else:
                logger.info("音声録音スレッドが終了しました。")
        self.audio_recording_thread = None

        # ボタン状態は stop_all_tasks で制御

        # 動画ファイルを保存
        if self.recorded_frames:
            now = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"recording_{now}.mp4"
            # ★★★ 保存パスは self.save_folder_name.get() から取得 ★★★
            save_folder = self.save_folder_name.get()
            if not save_folder or not os.path.isdir(save_folder):
                logger.error(
                    f"無効な保存フォルダです: {save_folder}。動画を保存できません。"
                )
                # 必要であればここでエラーメッセージ表示
                messagebox.showerror(
                    "保存エラー",
                    f"指定された保存フォルダが見つかりません:\n{save_folder}",
                )
            else:
                output_filepath = os.path.join(save_folder, output_filename)
                logger.info(f"動画ファイルの保存を開始します: {output_filepath}")

                # 保存処理を別スレッドで行う (UIが固まるのを防ぐため)
                save_thread = threading.Thread(
                    target=self._save_video_with_audio,  # この中でノート作成がトリガーされる
                    args=(output_filepath,),
                    name="VideoSaveThread",
                    daemon=True,
                )
                save_thread.start()
                logger.info("動画保存スレッドを開始しました。")

        else:
            logger.warning("録画フレームがないため、動画ファイルは保存されません。")

        self.recording_start_time = None
        self.update_recording_status()  # ステータスを更新

        # エラーチェック
        if self.error_occurred_in_recording_thread:
            messagebox.showerror(
                "録画エラー",
                "録画中にエラーが発生しました。詳細はログを確認してください。",
            )
            self.error_occurred_in_recording_thread = False  # フラグをリセット

        logger.info("録画停止処理を完了しました。")

    def start_note_creation(self, video_filepath):  # 引数に video_filepath を追加
        """指定された動画ファイルパスでノート作成処理を開始する"""
        # --- GUI操作制限 ---
        logger.info("ノート作成開始: GUI操作を制限します。")
        # ノート作成は stop_all_tasks 後に呼ばれるため、ボタンは既に無効化されているはずだが念のため
        self.start_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.DISABLED)  # stop_all_tasks で無効化される
        self.refresh_window_list_button.config(state=tk.DISABLED)
        self.window_listbox.config(state=tk.DISABLED)
        # 閉じるボタンを無効化
        self.root.protocol(
            "WM_DELETE_WINDOW",
            lambda: logger.warning("ノート作成中はウィンドウを閉じられません。"),
        )
        # --- ここまで ---
        if not self.gemini_client:
            self.note_creation_status.set("ノート作成不可: APIクライアント未設定")
            logger.warning(
                "Gemini API クライアントが設定されていないため、ノート作成を開始できません。"
            )
            return
        if not video_filepath or not os.path.exists(video_filepath):
            self.note_creation_status.set(
                f"ノート作成不可: 動画ファイルが見つかりません ({video_filepath})"
            )
            logger.error(f"指定された動画ファイルが見つかりません: {video_filepath}")
            return

        # ★ 開始時のステータスを設定
        self.note_creation_status.set("ノート生成準備中...")
        self.note_result_message.set("")
        logger.info(f"ノート作成処理を開始します。対象動画: {video_filepath}")

        # Gemini APIの呼び出しとmd生成の処理 (内部関数)
        def api_call_and_md_gen():
            try:
                prompt = """添付された動画の内容を分析し、各トピックごとに指定されたJSONスキーマに沿って情報を抽出してください。情報抽出の際は、次の条件を満たすようにしてください。\n
                - クライアントの理解レベルは「大学生程度」を想定し、専門的な表現は避け、平易な日本語で記述してください。\n
                - 専門用語は、一般的に意味が伝わらないと判断されるもののみ簡潔に解説を加えてくださ  い。\n
                - 応答は常に日本語で行い、入力が他言語であっても翻訳せず日本語で応答してください。\n
                - 動画内で本筋と無関係な内容（雑談・ノイズ等）は除外してください。\n
                - 各トピックの抽出・解釈・変換は、推論のプロセス（Thinking）と最終出力（Result）に分けて順序立てて記述してください。\n

                # Output Format\n
                指定されたJSONスキーマに従って、出力してください。\n

                # Notes\n

                - JSONスキーマの具体的な構造は別途提供される前提で処理を行ってください。\n
                - thinkingセクションでは、動画内の文脈・発話・論理の流れを明示的に示してください（単なる要約ではなく、どのようにその要約に至ったかを示してください）。\n
                - topicsセクションでは、思考の結論として整理された最終出力を記述してください。\n"""
                # スキーマ定義は変更なし (省略)
                schema = {
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "動画のタイトル",
                        },
                        "summary": {
                            "type": "string",
                            "description": "動画全体の内容の要約を3000字以内で",
                        },
                        "topics": {
                            "type": "array",
                            "description": "動画内の各トピックの情報",
                            "items": {
                                "type": "object",
                                "description": "個別トピックの情報",
                                "properties": {
                                    "topic_title": {
                                        "type": "string",
                                        "description": "トピックのタイトル",
                                    },
                                    "topic_keywords": {
                                        "type": "array",
                                        "description": "トピックのキーワード一覧",
                                        "items": {
                                            "type": "string",
                                            "description": "キーワード",
                                        },
                                    },
                                    "topic_summary": {
                                        "type": "string",
                                        "description": "トピックの要約を1000字以内で",
                                    },
                                    "topic_points": {
                                        "type": "array",
                                        "description": "トピックの重要ポイント一覧",
                                        "items": {
                                            "type": "string",
                                            "description": "重要ポイントを詳しく",
                                        },
                                    },
                                    "technical_term": {
                                        "type": "array",
                                        "description": "専門用語の解説一覧",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "word": {
                                                    "type": "string",
                                                    "description": "専門用語",
                                                },
                                                "explanation": {
                                                    "type": "string",
                                                    "description": "専門用語の解説",
                                                },
                                            },
                                            "required": ["word", "explanation"],
                                        },
                                    },
                                    "topic_thinking": {
                                        "type": "string",
                                        "description": "各トピックでの思考過程",
                                    },
                                },
                                "required": [
                                    "topic_title",
                                    "topic_summary",
                                    "topic_keywords",
                                    "topic_points",
                                    "topic_thinking",
                                ],
                                "propertyOrdering": [
                                    "topic_thinking",
                                    "topic_title",
                                    "topic_summary",
                                    "topic_keywords",
                                    "topic_points",
                                    "technical_term",
                                ],
                            },
                        },
                        "thinking": {
                            "type": "string",
                            "description": "全体の思考過程",
                        },
                    },
                    "required": ["title", "summary", "topics", "thinking"],
                    "propertyOrdering": ["thinking", "title", "summary", "topics"],
                }

                logger.info(f"動画ファイルをアップロード開始: {video_filepath}")
                # --- ステータス更新: アップロード中 ---
                self.root.after(
                    0, self.note_creation_status.set, "動画アップロード中..."
                )

                # --- ファイルアップロードと状態待機 (再試行ロジック付き) ---
                max_retries = 3
                retry_delay = 1  # seconds
                video_file = None
                upload_and_wait_success = False

                for attempt in range(max_retries):
                    try:
                        logger.info(
                            f"動画ファイルアップロード試行 {attempt + 1}/{max_retries}: {video_filepath}"
                        )
                        self.root.after(
                            0,
                            self.note_creation_status.set,
                            f"動画ファイルアップロード試行 {attempt + 1}/{max_retries}...",
                        )

                        # 1. ファイルアップロード
                        video_file = self.gemini_client.files.upload(
                            file=video_filepath
                        )
                        logger.info(
                            f"アップロード開始: {video_file.name}, State: {video_file.state}"
                        )
                        self.root.after(
                            0,
                            self.note_creation_status.set,
                            f"アップロード開始: {video_file.state}",
                        )

                        # 2. ファイルがACTIVEになるまで待機
                        polling_interval = 5
                        timeout_seconds = 1800  # 30分
                        start_poll_time = time.time()
                        while video_file.state != "ACTIVE":
                            if time.time() - start_poll_time > timeout_seconds:
                                raise TimeoutError(
                                    f"ファイル処理がタイムアウトしました ({timeout_seconds}秒): {video_file.name}"
                                )

                            # --- ステータス更新: 処理中 ---
                            elapsed_time = time.time() - start_poll_time
                            status_text = f"動画処理中... ({video_file.state}, {elapsed_time:.1f}秒)"
                            self.root.after(
                                0, self.note_creation_status.set, status_text
                            )
                            logger.info(status_text)

                            time.sleep(polling_interval)
                            video_file = self.gemini_client.files.get(
                                name=video_file.name
                            )  # 最新の状態を取得

                        logger.info(f"ファイルがACTIVEになりました: {video_file.name}")
                        self.root.after(
                            0,
                            self.note_creation_status.set,
                            "ファイル処理完了 (ACTIVE)",
                        )
                        upload_and_wait_success = True
                        break  # 成功したらループを抜ける

                    except TimeoutError as timeout_err:
                        logger.warning(
                            f"試行 {attempt + 1}/{max_retries}: ファイル処理タイムアウト - {timeout_err}"
                        )
                        # タイムアウトの場合はアップロードされたファイルを削除試行 (失敗しても無視)
                        if video_file and video_file.name:
                            try:
                                logger.info(
                                    f"タイムアウトしたファイル {video_file.name} を削除します。"
                                )
                                self.gemini_client.files.delete(name=video_file.name)
                            except Exception as delete_err:
                                logger.warning(
                                    f"タイムアウトしたファイルの削除中にエラー: {delete_err}"
                                )
                        video_file = None  # ファイル参照をリセット

                        if attempt < max_retries - 1:
                            logger.info(f"{retry_delay}秒後に再試行します...")
                            self.root.after(
                                0,
                                self.note_creation_status.set,
                                f"タイムアウト、再試行 ({attempt + 1}/{max_retries})... {retry_delay}秒待機",
                            )
                            time.sleep(retry_delay)
                            retry_delay *= 2  # 指数バックオフ
                        else:
                            logger.exception(
                                f"ファイル処理が最終的にタイムアウトしました (試行 {max_retries}回): {timeout_err}"
                            )
                            self.root.after(
                                0,
                                self.note_creation_status.set,
                                f"エラー: ファイル処理タイムアウト (試行 {max_retries}回)",
                            )
                            self.root.after(
                                0,
                                self.finish_note_creation,
                                False,
                                f"ファイル処理タイムアウト (試行 {max_retries}回): {timeout_err}",
                            )
                            return  # 処理中断

                    except Exception as upload_err:
                        logger.warning(
                            f"試行 {attempt + 1}/{max_retries}: アップロードまたは状態確認中にエラー - {upload_err}"
                        )
                        # エラー発生時もファイルを削除試行
                        if video_file and video_file.name:
                            try:
                                logger.info(
                                    f"エラーが発生したファイル {video_file.name} を削除します。"
                                )
                                self.gemini_client.files.delete(name=video_file.name)
                            except Exception as delete_err:
                                logger.warning(
                                    f"エラーファイルの削除中にエラー: {delete_err}"
                                )
                        video_file = None  # ファイル参照をリセット

                        if attempt < max_retries - 1:
                            logger.info(f"{retry_delay}秒後に再試行します...")
                            self.root.after(
                                0,
                                self.note_creation_status.set,
                                f"エラー発生、再試行 ({attempt + 1}/{max_retries})... {retry_delay}秒待機",
                            )
                            time.sleep(retry_delay)
                            retry_delay *= 2  # 指数バックオフ
                        else:
                            logger.exception(
                                f"ファイルアップロード/処理が最終的に失敗しました (試行 {max_retries}回): {upload_err}"
                            )
                            self.root.after(
                                0,
                                self.note_creation_status.set,
                                f"エラー: アップロード/処理失敗 (試行 {max_retries}回)",
                            )
                            self.root.after(
                                0,
                                self.finish_note_creation,
                                False,
                                f"アップロード/処理エラー (試行 {max_retries}回): {upload_err}",
                            )
                            return  # 処理中断

                # ループを抜けた後、成功フラグを確認
                if not upload_and_wait_success or video_file is None:
                    logger.error(
                        "不明な理由でファイルアップロード/処理に失敗しました。"
                    )
                    self.root.after(
                        0,
                        self.note_creation_status.set,
                        "エラー: ファイル処理失敗 (不明)",
                    )
                    self.root.after(
                        0, self.finish_note_creation, False, "ファイル処理エラー (不明)"
                    )
                    return

                # --- ステータス更新: ノート生成中 ---
                self.root.after(0, self.note_creation_status.set, "ノート生成中...")

                # Gemini APIに要約リクエストを送信
                logger.info("Gemini APIに要約リクエストを送信します...")
                # ★ モデル名を修正 (例: gemini-1.5-pro-latest)
                #    config ではなく generation_config を使用
                response = self.api_call_with_retry(video_file, prompt, schema)
                self.gemini_client.files.delete(name=video_file.name)
                summary_text = response.text
                logger.info("Gemini API から応答を取得しました。")
                # --- ステータス更新: 応答処理中 ---
                self.root.after(0, self.note_creation_status.set, "応答を処理中...")

                # mdファイル生成
                try:
                    summary_data = json.loads(summary_text)
                    logger.info("応答のJSONパースに成功しました。")

                    md_filename = f"note_{os.path.splitext(os.path.basename(video_filepath))[0]}.md"
                    md_filepath = os.path.join(
                        os.path.dirname(video_filepath), md_filename
                    )
                    logger.info(f"Markdownファイルを生成します: {md_filepath}")

                    markd = Markdown()
                    markd.add_header(summary_data.get("title", "タイトルなし"))
                    markd.add_header("全体要約", 2)
                    markd.add_text(summary_data.get("summary", "要約なし"))
                    markd.add_header("トピック詳細", 2)
                    topics = summary_data.get("topics", [])
                    if topics:
                        for i, topic in enumerate(topics):
                            topic_title = topic.get("topic_title", f"トピック {i+1}")
                            markd.add_header(topic_title, 3)
                            markd.add_linebreak()
                            keywords = topic.get("topic_keywords", [])
                            if keywords:
                                markd.add_text("キーワード:")
                                for kw in keywords:
                                    markd.add_list_item(f"{kw}")
                            markd.add_linebreak()
                            topic_summary = topic.get("topic_summary", "要約なし")
                            markd.add_text("要約:")
                            markd.add_text(topic_summary)
                            markd.add_linebreak()
                            points = topic.get("topic_points", [])
                            if points:
                                markd.add_text("ポイント:")
                                for pt in points:
                                    markd.add_list_item(f"{pt}")
                            terms = topic.get("technical_term", [])
                            if terms:
                                markd.add_text("専門用語:")
                                for term in terms:
                                    word = term.get("word", "")
                                    explanation = term.get("explanation", "")
                                    markd.add_list_item(f"{word} : {explanation}")
                    else:
                        markd.add_text("トピック情報はありません。")
                    # doc = Document()
                    # doc.add_heading(summary_data.get("title", "タイトルなし"), 0)
                    # doc.add_heading("全体要約", level=1)
                    # doc.add_paragraph(summary_data.get("summary", "要約なし"))
                    # doc.add_heading("トピック詳細", level=1)
                    # topics = summary_data.get("topics", [])
                    # if topics:
                    #     for i, topic in enumerate(topics):
                    #         topic_title = topic.get("topic_title", f"トピック {i+1}")
                    #         doc.add_heading(topic_title, level=2)
                    #         keywords = topic.get("topic_keywords", [])
                    #         if keywords:
                    #             doc.add_paragraph("キーワード:")
                    #             for kw in keywords:
                    #                 doc.add_paragraph(f"- {kw}", style="List Bullet")
                    #         topic_summary = topic.get("topic_summary", "要約なし")
                    #         doc.add_paragraph("要約:")
                    #         doc.add_paragraph(topic_summary)
                    #         points = topic.get("topic_points", [])
                    #         if points:
                    #             doc.add_paragraph("ポイント:")
                    #             for pt in points:
                    #                 doc.add_paragraph(f"- {pt}", style="List Bullet")
                    #         terms = topic.get("technical_term", [])
                    #         if terms:
                    #             doc.add_paragraph("専門用語:")
                    #             for term in terms:
                    #                 md = term.get("md", "")
                    #                 explanation = term.get("explanation", "")
                    #                 doc.add_paragraph(
                    #                     f"- {md} : {explanation}", style="List Bullet"
                    #                 )
                    #         doc.add_paragraph()
                    # else:
                    #     doc.add_paragraph("トピック情報はありません。")

                    directory = os.path.dirname(md_filepath)
                    if not os.path.exists(directory):
                        os.makedirs(directory, exist_ok=True)
                    markd.save(md_filepath)
                    logger.info(f"mdファイルを保存しました: {md_filepath}")
                    # UIスレッドでステータスを更新 (成功)
                    self.root.after(0, self.finish_note_creation, True, md_filepath)

                except json.JSONDecodeError as json_err:
                    logger.error(f"Gemini応答JSON解析失敗: {json_err}")
                    logger.error(f"受信テキスト(一部): {summary_text[:500]}...")
                    error_msg = f"API応答JSON解析エラー: {json_err}\n応答(一部): {summary_text[:200]}..."
                    self.root.after(0, self.finish_note_creation, False, error_msg)
                except Exception as md_err:
                    logger.exception(f"mdファイル生成エラー: {md_err}")
                    self.root.after(
                        0,
                        self.finish_note_creation,
                        False,
                        f"md生成エラー: {md_err}",
                    )

            except TimeoutError as te:
                logger.error(f"ノート作成タイムアウト: {te}")
                self.root.after(
                    0,
                    self.finish_note_creation,
                    False,
                    f"ファイル処理タイムアウト: {te}",
                )
            except Exception as e:
                logger.exception("ノート作成処理中エラー")
                # ★ ファイル削除処理を追加
                if "video_file" in locals() and video_file:
                    try:
                        logger.info(
                            f"エラー発生のためアップロードファイル削除: {video_file.name}"
                        )
                        self.gemini_client.files.delete(name=video_file.name)
                        logger.info("ファイル削除成功")
                    except Exception as delete_err:
                        logger.error(f"アップロードファイル削除失敗: {delete_err}")
                # UIスレッドでステータスを更新 (失敗)
                self.root.after(0, self.finish_note_creation, False, str(e))
            # ★ finally ブロックを削除 (エラーハンドリング内でファイル削除を行うため)

        # 別スレッドで実行
        note_thread = threading.Thread(
            target=api_call_and_md_gen, name="NoteCreationThread", daemon=True
        )
        note_thread.start()

    def api_call_with_retry(
        self, video_file, prompt, schema, max_retries=3, retry_delay=5
    ):
        """リトライロジックを持つAPIコール"""
        for attempt in range(max_retries):
            try:
                logger.info(f"Gemini API呼び出し試行 {attempt+1}/{max_retries}")
                # 選択されたモデル名を取得
                selected_model_name = self.selected_gemini_model.get()
                logger.info(f"使用する Gemini モデル: {selected_model_name}")
                response = self.gemini_client.models.generate_content(
                    model=selected_model_name,  # 選択されたモデルを使用
                    contents=[video_file, prompt],
                    config={
                        "response_mime_type": "application/json",
                        "response_schema": schema,
                    },
                )
                return response
            except google.genai.errors.ServerError as e:
                logger.warning(f"APIサーバーエラー（試行 {attempt+1}）: {e}")
                if attempt < max_retries - 1:
                    logger.info(f"{retry_delay}秒後に再試行します...")
                    time.sleep(retry_delay)
                else:
                    logger.error(f"最大試行回数に達しました。エラー: {e}")
                    raise

    def finish_note_creation(self, success, result_path_or_error):
        """ノート作成処理の完了をUIに反映する"""
        if success:
            # ★ 成功時のステータス
            self.note_creation_status.set("ノート生成完了")
            self.note_result_message.set(
                f"保存先: {os.path.basename(result_path_or_error)}"
            )  # ファイル名のみ表示
            logger.info(f"ノート作成成功: {result_path_or_error}")
            messagebox.showinfo(
                "ノート作成完了",
                f"ノートが正常に作成されました。\n{result_path_or_error}",
            )
        else:
            # ★ 失敗時のステータス
            self.note_creation_status.set("ノート生成失敗")
            # エラーメッセージを短縮して表示
            error_short = str(result_path_or_error).split("\n")[0][:100] + "..."
            self.note_result_message.set(
                f"エラー: {error_short}"
            )  # 修正: error_short を使用
            logger.error(f"ノート作成失敗: {result_path_or_error}")
            messagebox.showerror(
                "ノート作成エラー",
                f"ノートの作成中にエラーが発生しました:\n{result_path_or_error}",
            )

        # --- GUI操作制限解除 ---
        logger.info("ノート作成完了/失敗: GUI操作制限を解除します。")
        self.start_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)  # 停止ボタンは常に無効で良い
        self.folder_entry.config(state=tk.NORMAL)
        self.refresh_window_list_button.config(state=tk.NORMAL)
        self.window_listbox.config(state=tk.NORMAL)
        self.model_combobox.config(state="readonly")
        self.root.update()
        if hasattr(self, "original_on_closing"):
            self.root.protocol("WM_DELETE_WINDOW", self.original_on_closing)
        else:
            self.root.protocol("WM_DELETE_WINDOW", self.root.destroy)

    def stop_all_tasks(self):
        """全てのキャプチャ・録画タスクを停止する"""
        logger.info("全てのタスクの停止処理を開始します。")

        # ★★★ 修正点: 停止ボタンをすぐに無効化 ★★★
        self.stop_button.config(state=tk.DISABLED)
        # 開始ボタンは、関連処理がすべて完了するまで無効のままにする
        # (finish_note_creation またはこの関数の最後で有効に戻す)

        was_recording = self.is_recording  # 録画中だったか記録

        # 録画停止 (音声も内部で停止される)
        if self.is_recording:
            self.stop_recording()  # stop_recording が呼ばれる (内部で保存とノート作成トリガー)

        # スクリーンショット停止
        if self.is_capturing_screenshot:
            self.stop_screenshot_capture()

        if was_recording:
            logger.info(
                "録画を停止しました。動画保存とノート作成がバックグラウンドで実行されます（成功した場合）。"
            )
            # この場合、UIの有効化は finish_note_creation に任せる
            # finish_note_creation が呼ばれないケース(動画保存失敗など)も考慮すると、
            # ここで start_button などを有効にすべきではない。
        else:
            logger.info("スクリーンショットキャプチャを停止しました。")
            # スクリーンショットのみ停止した場合、ノート作成は走らないので、
            # ここでUIを操作可能に戻す。
            self.start_button.config(state=tk.NORMAL)
            self.folder_entry.config(state=tk.NORMAL)
            self.refresh_window_list_button.config(state=tk.NORMAL)
            self.window_listbox.config(state=tk.NORMAL)
            self.model_combobox.config(state="readonly")
            # 閉じるボタンの挙動も元に戻す (ノート作成がない場合)
            if hasattr(self, "original_on_closing"):
                self.root.protocol("WM_DELETE_WINDOW", self.original_on_closing)
            else:
                self.root.protocol("WM_DELETE_WINDOW", self.root.destroy)

        logger.info("全てのタスクの停止処理を完了しました。")

    def prepare_save_folder(self):
        """保存フォルダの準備（作成、権限チェック）を行う"""
        initial_folder_name_from_ui = self.save_folder_name.get()
        current_working_directory = os.getcwd()  # ★ 現在の作業ディレクトリをログに出力
        logger.info(
            f"prepare_save_folder - UIから取得した初期フォルダ名: '{initial_folder_name_from_ui}'"
        )
        logger.info(
            f"prepare_save_folder - 現在の作業ディレクトリ: {current_working_directory}"
        )

        folder_name = initial_folder_name_from_ui
        if not folder_name:
            # フォルダ名が空の場合、現在時刻でデフォルト名を生成
            folder_name = f"capture_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            # self.save_folder_name.set(folder_name) # UIへのセットは成功後の方が良いかもしれない
            logger.info(
                f"prepare_save_folder - 保存フォルダ名が指定されなかったため、デフォルト名を生成しました: {folder_name}"
            )

        # folder_name が絶対パスか相対パスかで処理を分ける
        if os.path.isabs(folder_name):
            save_path = folder_name
            logger.info(
                f"prepare_save_folder - 入力されたフォルダ名は絶対パスとして扱います: {save_path}"
            )
        else:
            # 相対パスの場合、現在の作業ディレクトリを基準に絶対パスを生成
            save_path = os.path.join(current_working_directory, folder_name)
            # os.path.abspath は上記とほぼ同等だが、より明示的に join を使う
            # save_path = os.path.abspath(folder_name)
            logger.info(
                f"prepare_save_folder - 相対パスとして解釈し、絶対パスに変換しました (基準: {current_working_directory}): {save_path}"
            )

        logger.info(
            f"prepare_save_folder - 最終的な保存先フォルダ (os.makedirs対象): {save_path}"
        )

        try:
            # フォルダが存在しない場合は作成
            if not os.path.exists(save_path):
                logger.info(
                    f"prepare_save_folder - os.makedirs を呼び出します: Path='{save_path}', exist_ok=True"
                )
                os.makedirs(save_path, exist_ok=True)
                # makedirs がエラーを出さなかった場合、通常は作成されているはず
                logger.info(
                    f"prepare_save_folder - フォルダを作成しました (または既に存在していました): {save_path}"
                )
                if os.path.exists(save_path):
                    logger.info(
                        f"prepare_save_folder - フォルダ作成後の存在確認OK: {save_path}"
                    )
                else:
                    # このケースは稀だが、makedirsが成功したように見えても実際には作成されていない場合
                    logger.error(
                        f"prepare_save_folder - フォルダ作成後に存在確認NG (os.makedirsはエラーなし): {save_path}"
                    )
            else:
                logger.info(
                    f"prepare_save_folder - フォルダは既に存在します: {save_path}"
                )

            # 書き込み権限をチェック (簡易的な方法)
            test_file_path = os.path.join(save_path, ".permission_test")
            logger.info(
                f"prepare_save_folder - 書き込み権限テストファイルを作成します: {test_file_path}"
            )
            with open(test_file_path, "w") as f:
                f.write("test")
            os.remove(test_file_path)
            logger.info(
                f"prepare_save_folder - フォルダへの書き込み権限を確認しました: {save_path}"
            )
            self.save_folder_name.set(
                save_path
            )  # フォルダ準備成功後にUIに絶対パスを反映
            return True

        except OSError as e:
            # ★ OSError の詳細情報をログに出力
            logger.exception(
                f"prepare_save_folder - 保存フォルダの準備中にOSErrorが発生しました: Path='{save_path}', ErrorNo={e.errno}, ErrorMsg='{e.strerror}', Details='{e}'"
            )
            messagebox.showerror(
                "フォルダエラー",
                f"保存フォルダの作成またはアクセスに失敗しました:\n{save_path}\n\nエラーコード: {e.errno}\n詳細: {e.strerror}",
            )
            return False
        except Exception as e:
            logger.exception(
                f"prepare_save_folder - 保存フォルダの準備中に予期せぬエラーが発生しました: Path='{save_path}', Error='{e}'"
            )
            messagebox.showerror(
                "フォルダエラー",
                f"保存フォルダの準備中に予期せぬエラーが発生しました:\n{save_path}\n\nエラー詳細:\n{e}",
            )
            return False

    # --- ステータス更新メソッド ---
    def update_screenshot_status(self):
        """スクリーンショットのステータス表示を更新する"""
        if self.is_capturing_screenshot:
            status_text = f"実行中... ({self.screenshot_saved_count}枚保存)"
            if self.last_saved_screenshot_filename:
                status_text += f"\n最終保存: {os.path.basename(self.last_saved_screenshot_filename)}"
            self.screenshot_status_label.config(text=status_text)
            # 1秒後に再度更新
            self.root.after(1000, self.update_screenshot_status)
        else:
            status_text = "停止中"
            if self.last_saved_screenshot_filename:
                status_text += f" (最終保存: {os.path.basename(self.last_saved_screenshot_filename)})"
            elif self.screenshot_saved_count > 0:
                status_text += f" ({self.screenshot_saved_count}枚保存)"
            else:
                status_text = "待機中..."
            self.screenshot_status_label.config(text=status_text)

    def update_recording_status(self):
        """録画のステータス表示を更新する"""
        if self.is_recording and self.recording_start_time:
            elapsed_seconds = int(time.time() - self.recording_start_time)
            minutes, seconds = divmod(elapsed_seconds, 60)
            status_text = f"録画中... {minutes:02d}:{seconds:02d}"
            if self.last_saved_video_filename:  # 保存が完了していれば表示
                status_text += (
                    f"\n最終保存: {os.path.basename(self.last_saved_video_filename)}"
                )
            self.recording_status_label.config(text=status_text)
            # 1秒後に再度更新
            self.root.after(1000, self.update_recording_status)
        else:
            status_text = "停止中"
            if self.last_saved_video_filename:
                status_text += (
                    f" (最終保存: {os.path.basename(self.last_saved_video_filename)})"
                )
            else:
                status_text = "待機中..."
            self.recording_status_label.config(text=status_text)

    # --- スクリーンショット関連メソッド ---
    def start_screenshot_capture(self):
        """スクリーンショットの連続キャプチャを開始する"""
        if self.is_capturing_screenshot:
            return

        hwnd = self.selected_window_handle.get()  # 選択されたウィンドウハンドルを取得

        logger.info(
            f"スクリーンショットキャプチャを開始します。対象ウィンドウハンドル: {hwnd if hwnd != 0 else 'プライマリモニター'}"
        )
        # フォルダ準備は start_tasks で行う

        # スクリーンショットキャプチャスレッドを開始 (hwnd を渡す)
        self.screenshot_thread = threading.Thread(
            target=self.screenshot_capture_loop,
            args=(hwnd,),
            name="ScreenshotThread",
            daemon=True,
        )
        self.is_capturing_screenshot = True
        self.screenshot_saved_count = 0  # カウンタリセット
        self.last_saved_screenshot_filename = ""  # ファイル名リセット
        self.screenshot_thread.start()

        # self.start_screenshot_button.config(state=tk.DISABLED) # 個別ボタン削除
        # self.stop_screenshot_button.config(state=tk.NORMAL) # 個別ボタン削除
        # ボタン状態は start_tasks / stop_all_tasks で制御
        self.update_screenshot_status()  # ステータス更新開始

    def stop_screenshot_capture(self):
        """スクリーンショットの連続キャプチャを停止する"""
        if not self.is_capturing_screenshot:
            return

        logger.info("スクリーンショットキャプチャの停止処理を開始します。")
        self.is_capturing_screenshot = False
        if self.screenshot_thread and self.screenshot_thread.is_alive():
            logger.info("スクリーンショットキャプチャスレッドの終了を待機します...")
            self.screenshot_thread.join(timeout=5.0)  # タイムアウトを設定
            if self.screenshot_thread.is_alive():
                logger.warning(
                    "スクリーンショットキャプチャスレッドが時間内に終了しませんでした。"
                )
            else:
                logger.info("スクリーンショットキャプチャスレッドが終了しました。")
        self.screenshot_thread = None

        # self.start_screenshot_button.config(state=tk.NORMAL) # 個別ボタン削除
        # self.stop_screenshot_button.config(state=tk.DISABLED) # 個別ボタン削除
        # ボタン状態は start_tasks / stop_all_tasks で制御
        self.update_screenshot_status()  # ステータス更新

        # エラーチェック
        if self.error_occurred_in_screenshot_thread:
            messagebox.showerror(
                "スクリーンショットエラー",
                "スクリーンショット取得中にエラーが発生しました。詳細はログを確認してください。",
            )
            self.error_occurred_in_screenshot_thread = False  # フラグをリセット

        logger.info("スクリーンショットキャプチャの停止処理を完了しました。")

    def screenshot_capture_loop(self, hwnd):  # 引数 hwnd を追加
        """指定されたウィンドウまたは画面全体のスクリーンショットを定期的に取得し、変化があれば保存するループ"""
        logger.info(
            f"スクリーンショットキャプチャループを開始します。対象ウィンドウハンドル: {hwnd if hwnd != 0 else 'プライマリモニター'}"
        )
        interval_seconds = 1  # 取得間隔（秒）
        self.last_screenshot_image = None  # 前回の画像をリセット

        try:  # sct オブジェクトの初期化を try の外に出す
            sct = mss.mss()
        except Exception as e:
            logger.exception(f"mss の初期化中にエラー: {e}")
            self.error_occurred_in_screenshot_thread = True
            return  # 初期化失敗時はループに入らない

        while self.is_capturing_screenshot:
            start_capture_time = time.time()
            try:
                # キャプチャ領域を決定
                monitor = None
                current_hwnd = hwnd  # ループ内で hwnd を変更する可能性があるのでコピー
                if current_hwnd != 0:
                    try:
                        # ウィンドウが存在するか確認
                        if not win32gui.IsWindow(current_hwnd):
                            logger.warning(
                                f"ウィンドウハンドル {current_hwnd} が無効です。プライマリモニターを対象にします。"
                            )
                            current_hwnd = (
                                0  # 無効ならプライマリモニターにフォールバック
                            )
                        else:
                            window_rect = win32gui.GetWindowRect(current_hwnd)
                            left, top, right, bottom = window_rect
                            width = right - left
                            height = bottom - top
                            if width > 0 and height > 0:
                                monitor = {
                                    "left": left,
                                    "top": top,
                                    "width": width,
                                    "height": height,
                                }
                                # logger.debug(f"対象ウィンドウ領域: {monitor}") # デバッグ用
                            else:
                                logger.warning(
                                    f"ウィンドウサイズが無効です ({width}x{height})。プライマリモニターを対象にします。"
                                )
                                current_hwnd = 0  # ウィンドウが無効ならプライマリモニターにフォールバック
                    except Exception as e:
                        logger.warning(
                            f"ウィンドウ情報の取得に失敗しました: {e}。プライマリモニターを対象にします。"
                        )
                        current_hwnd = 0  # エラー時もプライマリモニターにフォールバック

                if (
                    monitor is None
                ):  # current_hwnd が 0 またはウィンドウ情報取得失敗の場合
                    # プライマリモニターを取得 (monitors[0] は全画面結合の場合があるので注意、通常は monitors[1])
                    # 以前の実装 sct.monitors[0] に合わせるか、録画処理 sct.monitors[1] に合わせるか
                    # ここでは monitors[1] を試す (より一般的)
                    if len(sct.monitors) > 1:
                        monitor = sct.monitors[1]
                    else:
                        monitor = sct.monitors[0]  # モニターが1つしかない場合
                    # logger.debug(f"対象領域: プライマリモニター {monitor}") # デバッグ用

                # スクリーンショットを取得
                screenshot = sct.grab(monitor)
                # Pillowイメージに変換
                current_image_pil = Image.frombytes(
                    "RGB", screenshot.size, screenshot.rgb
                )
                # OpenCV形式に変換 (比較用)
                current_image_cv = cv2.cvtColor(
                    np.array(current_image_pil), cv2.COLOR_RGB2BGR
                )

                if self.last_screenshot_image is not None:
                    # 前回と比較して変化があるか確認
                    if not self.is_similar(
                        current_image_cv, self.last_screenshot_image
                    ):
                        logger.info(
                            "画面に変化を検出しました。スクリーンショットを保存します。"
                        )
                        if self.save_screenshot_image(
                            current_image_cv
                        ):  # 保存処理を呼び出し
                            self.screenshot_saved_count += 1
                    # else:
                    #     logger.debug("画面に変化はありません。") # デバッグ用
                else:
                    # 最初のスクリーンショットは必ず保存
                    logger.info("最初のスクリーンショットを保存します。")
                    if self.save_screenshot_image(current_image_cv):
                        self.screenshot_saved_count += 1

                # 今回の画像を次回比較用に保持
                self.last_screenshot_image = current_image_cv

            except UnidentifiedImageError as e:
                logger.error(f"スクリーンショット画像の形式が認識できませんでした: {e}")
                self.error_occurred_in_screenshot_thread = True
                # エラーが発生してもループは継続するかもしれないが、一旦フラグを立てる
            except mss.ScreenShotError as sct_err:  # mss のエラーもキャッチ
                logger.error(
                    f"スクリーンショット取得中にエラーが発生しました: {sct_err}"
                )
                self.error_occurred_in_screenshot_thread = True
            except Exception as e:
                logger.exception(
                    f"スクリーンショット取得/比較中にエラーが発生しました: {e}"
                )
                self.error_occurred_in_screenshot_thread = True
                # エラーが発生してもループは継続するかもしれないが、一旦フラグを立てる

            # 次のキャプチャまでの待機時間
            elapsed_time = time.time() - start_capture_time
            sleep_time = max(0, interval_seconds - elapsed_time)
            if sleep_time > 0:
                time.sleep(sleep_time)

        logger.info("スクリーンショットキャプチャループを終了します。")

    def is_similar(self, img1_cv, img2_cv, threshold=0.83):
        """2つの画像の類似度を計算する (差分ベースの簡易比較)"""
        try:
            # グレースケールに変換
            gray1 = cv2.cvtColor(img1_cv, cv2.COLOR_BGR2GRAY)
            gray2 = cv2.cvtColor(img2_cv, cv2.COLOR_BGR2GRAY)

            # サイズが異なる場合はリサイズ (小さい方に合わせるか、固定サイズにする)
            if gray1.shape != gray2.shape:
                # 例: img1 のサイズに合わせる
                h, w = gray1.shape
                gray2 = cv2.resize(gray2, (w, h))
                logger.warning("比較画像のサイズが異なるためリサイズしました。")

            # 差分を計算
            diff = cv2.absdiff(gray1, gray2)

            # 差分が閾値以下のピクセルの割合を計算
            non_zero_count = np.count_nonzero(diff > 10)  # わずかな違いは無視
            total_pixels = diff.shape[0] * diff.shape[1]
            similarity = 1.0 - (non_zero_count / total_pixels)

            # logger.debug(f"画像類似度: {similarity:.4f}") # デバッグ用

            return similarity >= threshold
        except cv2.error as e:
            logger.error(f"画像比較 (is_similar) 中にOpenCVエラーが発生しました: {e}")
            return False  # エラー時は類似していないと判断
        except Exception as e:
            logger.exception(
                f"画像比較 (is_similar) 中に予期せぬエラーが発生しました: {e}"
            )
            return False  # エラー時は類似していないと判断

    def save_screenshot_image(self, image_cv):  # save_image から変更
        """スクリーンショット画像をファイルに保存する。エラー発生時はログに記録。"""
        save_folder = self.save_folder_name.get()
        if not save_folder:
            logger.error(
                "保存フォルダが設定されていません。スクリーンショットを保存できません。"
            )
            return False

        try:
            now = datetime.now()
            # ミリ秒を含むファイル名
            filename = f"screenshot_{now.strftime('%Y%m%d_%H%M%S')}_{now.microsecond // 1000:03d}.png"
            filepath = os.path.join(save_folder, filename)

            # OpenCV形式(BGR)の画像をPNGで保存
            success = cv2.imwrite(filepath, image_cv)

            if success:
                logger.info(f"スクリーンショットを保存しました: {filepath}")
                self.last_saved_screenshot_filename = (
                    filepath  # 最後に保存したファイル名を更新
                )
                return True
            else:
                logger.error(
                    f"スクリーンショットの保存に失敗しました (cv2.imwriteがFalseを返しました): {filepath}"
                )
                # cv2.imwrite が False を返す具体的な原因は特定しにくい場合がある
                # ディスク容量、権限、ファイルパスの問題などが考えられる
                return False
        except cv2.error as e:
            logger.error(
                f"スクリーンショット保存中にOpenCVエラーが発生しました: {e}. ファイルパス: {filepath if 'filepath' in locals() else 'N/A'}"
            )
            return False
        except Exception as e:
            logger.exception(
                f"スクリーンショット保存中に予期せぬエラーが発生しました: {e}. ファイルパス: {filepath if 'filepath' in locals() else 'N/A'}"
            )
            return False

    def on_closing(self):
        """ウィンドウが閉じられるときの処理"""
        logger.info("アプリケーション終了処理を開始します。")
        if self.is_capturing_screenshot or self.is_recording:
            if messagebox.askokcancel(
                "確認", "キャプチャまたは録画が実行中です。本当に終了しますか？"
            ):
                self.stop_all_tasks()  # 実行中のタスクを停止
                # stop_all_tasks内で録画停止→保存→ノート作成が走る可能性があるため、少し待つか、
                # ノート作成完了まで待つ仕組みが必要かもしれない。
                # 現状では、ノート作成が完了する前にウィンドウが閉じる可能性がある。
                logger.info("実行中のタスクを停止しました。")
                self.root.destroy()
            else:
                logger.info("アプリケーションの終了をキャンセルしました。")
                return  # 終了をキャンセル
        else:
            self.root.destroy()
        logger.info("アプリケーションを終了しました。")


if __name__ == "__main__":
    try:
        try:
            from ctypes import windll

            # Per-Monitor DPI Aware V2 (Windows 10 Creators Update以降)
            # windll.shcore.SetProcessDpiAwareness(2)
            # System DPI Aware (より古いWindowsや互換性重視の場合)
            # windll.shcore.SetProcessDpiAwareness(1) # 既存のコードではこちらが使われている可能性

            # ★★★ 変更箇所 スタート ★★★
            # Per-Monitor DPI Aware V2 に設定してみる
            # これにより、各モニターのDPI設定が個別に認識されるようになる
            PROCESS_PER_MONITOR_DPI_AWARE = 2
            windll.shcore.SetProcessDpiAwareness(PROCESS_PER_MONITOR_DPI_AWARE)
            logger.info("DPI Awareness を Per-Monitor V2 (2) に設定しました。")
            # ★★★ 変更箇所 エンド ★★★

        except ImportError:
            logger.info(
                "ctypesモジュールが見つからないため、DPI Awareness は設定されませんでした (Windows以外の環境の可能性)。"
            )
        except AttributeError:
            # Per-Monitor V2 が利用できない古いWindowsの場合、System Aware を試す
            try:
                PROCESS_SYSTEM_DPI_AWARE = 1
                windll.shcore.SetProcessDpiAwareness(PROCESS_SYSTEM_DPI_AWARE)
                logger.info(
                    "DPI Awareness を System Aware (1) に設定しました (Per-Monitor V2 不可のため)。"
                )
            except AttributeError:
                logger.warning(
                    "DPI Awareness の設定に失敗しました (古いWindowsバージョンの可能性)。"
                )
            except Exception as e_sys:
                logger.warning(f"System DPI Awareness 設定中にエラー: {e_sys}")
        except Exception as e:
            logger.warning(f"DPI Awareness 設定中に予期せぬエラー: {e}")

        root = tk.Tk()
        app = SlideCaptureApp(root)
        root.mainloop()
    except Exception as e:
        # GUI初期化前などのエラーをキャッチ
        logger.critical(
            f"アプリケーションの起動中に致命的なエラーが発生しました: {e}",
            exc_info=True,
        )
        # コンソールにもエラーを表示
        print(f"致命的なエラーが発生しました: {e}")
        traceback.print_exc()
        # 簡単なメッセージボックスでユーザーに通知 (Tkinterが使える場合)
        try:
            error_root = tk.Tk()
            error_root.withdraw()  # メインウィンドウは表示しない
            messagebox.showerror(
                "起動エラー",
                f"アプリケーションの起動に失敗しました。\n詳細はログファイルを確認してください。\n\n{e}",
            )
            error_root.destroy()
        except tk.TclError:
            print(
                "メッセージボックスを表示できませんでした。"
            )  # Tkinterが初期化失敗した場合
        except Exception as msg_e:
            print(f"エラーメッセージ表示中にさらにエラー: {msg_e}")
