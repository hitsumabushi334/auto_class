# -*- coding: utf-8 -*-
"""
Gemini API を使った動画要約とWordファイル生成のテスト用スクリプト
main.pyから抽出した機能を単独で実行できるようにしています
"""

import os
import time
import logging
import json
import argparse
import traceback
from datetime import datetime
from docx import Document
from google import genai
from dotenv import load_dotenv

# --- ロギング設定 ---
log_filename = "gemini_video_summary_test.log"
log_format = "%(asctime)s - %(levelname)s - %(message)s"
logging.basicConfig(
    level=logging.INFO,
    format=log_format,
    handlers=[
        logging.FileHandler(log_filename, encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


def initialize_gemini_client():
    """Gemini APIクライアントを初期化する"""
    try:
        load_dotenv()  # .env ファイルを読み込む
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            logger.error("環境変数 GEMINI_API_KEY が設定されていません。")
            print("環境変数 GEMINI_API_KEY が設定されていません。")
            return None

        gemini_client = genai.Client(api_key=api_key)
        logger.info("Gemini API クライアントを設定しました。")
        return gemini_client

    except Exception as e:
        logger.exception("Gemini API の設定中にエラーが発生しました。")
        print(f"Gemini API の設定に失敗しました: {e}")
        return None


def generate_video_summary(gemini_client, video_filepath):
    """Gemini APIを使用して動画の要約を生成する"""
    if not gemini_client:
        logger.error("Gemini APIクライアントが初期化されていません。")
        return None

    if not os.path.exists(video_filepath):
        logger.error(f"指定された動画ファイルが見つかりません: {video_filepath}")
        return None

    try:
        prompt = "添付した動画の各トピックについて指定されたJSONスキーマに応じて内容をわかりやすく抽出してください。また、回答のレベルはクライアントが大学生レベルであることに注意して調節してください。応答は入力の言語の種類にかかわらず必ず日本語で行ってください。専門用語は一般的に意味が伝わらないと判断されるもののみを解説してください。"
        schema = {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "動画のタイトル",
                },
                "summary": {
                    "type": "string",
                    "description": "動画全体の内容の要約を3000字以内で",
                },
                "topics": {
                    "type": "array",
                    "description": "動画内の各トピックの情報",
                    "items": {
                        "type": "object",
                        "description": "個別トピックの情報",
                        "properties": {
                            "topic_title": {
                                "type": "string",
                                "description": "トピックのタイトル",
                            },
                            "topic_keyWords": {
                                "type": "array",
                                "description": "トピックのキーワード一覧",
                                "items": {
                                    "type": "string",
                                    "description": "キーワード",
                                },
                            },
                            "topic_summary": {
                                "type": "string",
                                "description": "トピックの要約を300字以内で",
                            },
                            "topic_points": {
                                "type": "array",
                                "description": "トピックの重要ポイント一覧",
                                "items": {
                                    "type": "string",
                                    "description": "重要ポイント",
                                },
                            },
                            "technical_term": {
                                "type": "array",
                                "description": "専門用語の解説一覧",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "word": {
                                            "type": "string",
                                            "description": "専門用語",
                                        },
                                        "explanation": {
                                            "type": "string",
                                            "description": "専門用語の解説",
                                        },
                                    },
                                    "required": ["word", "explanation"],
                                },
                            },
                        },
                        "required": [
                            "topic_title",
                            "topic_summary",
                            "topic_keyWords",
                            "topic_points",
                        ],
                    },
                },
            },
            "required": ["title", "summary", "topics"],
        }

        logger.info(f"動画ファイルをアップロード開始: {video_filepath}")
        video_file = gemini_client.files.upload(file=video_filepath)
        logger.info(
            f"動画ファイルをアップロード完了: {video_file.name}, State: {video_file.state}"
        )

        # ファイルがACTIVEになるまで待機
        polling_interval = 5  # ポーリング間隔 (秒)
        timeout_seconds = 300  # タイムアウト (秒)
        start_poll_time = time.time()

        while video_file.state != "ACTIVE":
            if time.time() - start_poll_time > timeout_seconds:
                logger.error(
                    f"ファイル処理がタイムアウトしました ({timeout_seconds}秒): {video_file.name}"
                )
                raise TimeoutError(
                    f"ファイル処理がタイムアウトしました ({timeout_seconds}秒): {video_file.name}"
                )

            logger.info(
                f"ファイル処理待機中... State: {video_file.state} (経過: {time.time() - start_poll_time:.1f}秒)"
            )
            time.sleep(polling_interval)
            video_file = gemini_client.files.get(name=video_file.name)  # 状態を再取得

        logger.info(f"ファイルがACTIVEになりました: {video_file.name}")

        # Gemini APIに要約リクエストを送信
        logger.info("Gemini APIに要約リクエストを送信します...")
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash-preview-04-17",  # 使用するモデル名は変更される可能性があります
            contents=[
                video_file,
                prompt,
            ],
            config={
                "response_mime_type": "application/json",
                "response_schema": schema,
            },
        )

        summary_text = response.text
        logger.info("Gemini API から応答を取得しました。")

        try:
            summary_data = json.loads(summary_text)
            logger.info("応答のJSONパースに成功しました。")
            return summary_data
        except json.JSONDecodeError as e:
            logger.error(f"Geminiからの応答JSONの解析に失敗しました: {e}")
            logger.error(
                f"受信したテキスト: {summary_text[:500]}..."
            )  # 最初の500文字を表示
            return None

    except TimeoutError as te:
        logger.error(f"タイムアウトエラー: {te}")
        return None
    except Exception as e:
        logger.exception(f"動画要約処理中にエラーが発生しました: {e}")
        return None


def create_word_document(summary_data, output_filepath):
    """要約データからWordドキュメントを生成する"""
    if not summary_data:
        logger.error("要約データがありません。Wordファイルを生成できません。")
        return False

    try:
        doc = Document()

        # タイトルを追加
        doc.add_heading(summary_data.get("title", "タイトルなし"), 0)

        # 全体要約を追加
        doc.add_heading("全体要約", level=1)
        doc.add_paragraph(summary_data.get("summary", "要約なし"))

        # 各トピックを追加
        doc.add_heading("トピック詳細", level=1)
        topics = summary_data.get("topics", [])

        if topics:
            for i, topic in enumerate(topics):
                # トピックタイトル（必須項目）
                topic_title = topic.get("topic_title", f"トピック {i+1}")
                doc.add_heading(topic_title, level=2)

                # キーワード
                keywords = topic.get("topic_keyWords", [])
                if keywords and len(keywords) > 0:
                    doc.add_paragraph("キーワード:")
                    for kw in keywords:
                        doc.add_paragraph(f"- {kw}", style="List Bullet")

                # トピック要約（必須項目）
                topic_summary = topic.get("topic_summary", "要約なし")
                doc.add_paragraph("要約:")
                doc.add_paragraph(topic_summary)

                # ポイント
                points = topic.get("topic_points", [])
                if points and len(points) > 0:
                    doc.add_paragraph("ポイント:")
                    for pt in points:
                        doc.add_paragraph(f"- {pt}", style="List Bullet")

                # 専門用語
                terms = topic.get("technical_term", [])
                if terms and len(terms) > 0:
                    doc.add_paragraph("専門用語:")
                    for term in terms:
                        word = term.get("word", "")
                        explanation = term.get("explanation", "")
                        doc.add_paragraph(
                            f"- {word} : {explanation}", style="List Bullet"
                        )

                doc.add_paragraph()  # トピック間にスペース
        else:
            doc.add_paragraph("トピック情報はありません。")

        # ファイルを保存
        directory = os.path.dirname(output_filepath)
        if not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)

        doc.save(output_filepath)
        logger.info(f"Wordファイルを保存しました: {output_filepath}")
        return True

    except Exception as e:
        logger.exception(f"Wordファイル生成中にエラーが発生しました: {e}")
        return False


def main():
    # コマンドライン引数の解析
    parser = argparse.ArgumentParser(
        description="Gemini APIを使った動画要約とWordファイル生成"
    )
    parser.add_argument("video_path", help="要約する動画ファイルのパス")
    parser.add_argument(
        "--output",
        "-o",
        help="出力するWordファイルのパス (省略時は動画と同じディレクトリに保存)",
    )
    args = parser.parse_args()

    video_path = os.path.abspath(args.video_path)

    if not os.path.exists(video_path):
        print(f"エラー: 指定された動画ファイルが見つかりません: {video_path}")
        return

    # 出力パスの設定
    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        # デフォルトは動画と同じディレクトリ
        video_dir = os.path.dirname(video_path)
        video_name = os.path.splitext(os.path.basename(video_path))[0]
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(video_dir, f"note_{video_name}_{now}.docx")

    # Gemini APIクライアントの初期化
    gemini_client = initialize_gemini_client()
    if not gemini_client:
        print("Gemini APIクライアントの初期化に失敗しました。処理を中止します。")
        return

    print(f"動画ファイル: {video_path}")
    print(f"出力ファイル: {output_path}")
    print("処理を開始します...")

    # 動画要約の生成
    summary_data = generate_video_summary(gemini_client, video_path)
    if not summary_data:
        print("動画要約の生成に失敗しました。処理を中止します。")
        return

    # 要約データを表示（確認用）
    print("\n=== 生成された要約データ ===")
    print(f"タイトル: {summary_data.get('title', 'タイトルなし')}")
    print(f"要約文字数: {len(summary_data.get('summary', ''))}")
    print(f"トピック数: {len(summary_data.get('topics', []))}")

    # Wordドキュメントの生成
    if create_word_document(summary_data, output_path):
        print(f"\n要約Wordファイルを生成しました: {output_path}")
    else:
        print("\nWordファイルの生成に失敗しました。")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logger.critical(f"予期せぬエラーが発生しました: {e}", exc_info=True)
        print(f"エラーが発生しました: {e}")
        traceback.print_exc()
